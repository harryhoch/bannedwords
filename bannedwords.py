#!/usr/bin/python

from docx import Document
import os
import json
import glob
import pdfplumber
import argparse
import re


def is_in_whitelist(word: str, whitelist:list) -> bool:
    if whitelist is None:
        return None
    matches  = [f for f in whitelist if word in f]
    return matches


# find starting positions of all of the times w occurs in the query.
def find_matches(w : str, para: str) -> list:
    res = []
    index = 0
    while index < len(para):
        index = para.find(w,index)
        if index == -1: 
            break
        res.append(index)
        index +=len(w)
    return res

def is_whitelisted(banned: str, whitelist: list,m: tuple) -> bool:
    """
        note that we are only interested in tuples that have any items that are not whitelisted. 
        this is not working. regexs are failing me. 
        instead. just look for indices of the banned word. if the banned word only shows up when anything from the whitelist shows up,
        it's whitelisted. 
        try find iter..
    """
    (i,para) = m
    # find where the banned word shows up in the text
    banned_matches = set(find_matches(banned,para))
    # find possible whitelist words in the white list. Which of those start with the banned words?
    matches =  set()
    for wh in whitelist:
        if wh.startswith(banned):
            wh_matches = set(find_matches(wh,para))
            matches = matches.union(wh_matches)
            # if the white list occurs everywhere that the banned word occurs, it's whitelisted.
    return matches == banned_matches


    print("looking to see if "+banned + " is whitelisted.. in " +para)
    # we want to find if the word (or phrase) matches on a word boundary.
    pat  = "(?<!\\S)"+banned+"(!\\S)+"
    match = re.search(pat,para)
    print(" match is..."+str(match))
    return match is  None

def new_func(banned):
    pat = "\\S"+banned+"\\S"
    return pat

def  filter_whitelist(word: str, whitelists: list, matches : list) -> list:
    """
    filter out  any possible instance of word that is matched by one of the
    items in the whitelist. 

    Args: 
        word (str): the banned word in question
        whitelists (list): list of matching whitelisted terms
        matches (list): a list of matches. each is a tuple of a paragraph 
            number and the text of the paragraph
    Returns:
        list: a list of tuples in the same form as matches, without any that would have
        been cleared by the whitelist
    """
    return [m for m in matches if not is_whitelisted(word,whitelists,m)]

def search_word_paras(paras: list, word: str,whitelist:list) -> list:
    """
    Search for paragraphs containing a specific word.

    Args:
        paras (list): List of paragraphs to search.
        word (str): Word to search for.
        whitelist: words to keep in . treateed differently

    Returns:
        list: List of paragraphs containing the word.
    """
    whitelists = is_in_whitelist(word,whitelist)
    results = search_word_direct(paras,word)
    if whitelists is not None:
        results = filter_whitelist(word,whitelists,results)
    return results

def search_word_direct(paras:list, word: str) -> list:
    """
    Do a direct word search, in which case we're looking to see if word 
    is found in paras in any form - as a word or as a substring.

    Args:
        paras (list): List of paragraphs to search.
        word (str): Word to search for.
            
    Returns:
        list: List of paragraphs containing the word.
    """
    results = [] 
    for i in range(len(paras)):
        if word.lower() in paras[i]:
            results.append((i,paras[i]))
    return results

def get_word_paras(d:Document) -> list:
    """
    Get a list of paragraphs from a Document object.

    Args:
        d (Document): Document object to extract paragraphs from.

    Returns:
        list: List of paragraphs from the document.
    """
    paras = [p.text for p in d.paragraphs]
    paras = [p.lower() for p in paras if len(p) > 0]
    return paras

def search_doc_for_categories(d: Document, categories:dict,whitelist:list) -> dict:

    paras  =get_word_paras(d)
    return search_paras_for_categories(paras,categories,whitelist)

def search_paras_for_categories(paras: list, categories:dict,whitelist:list) -> dict:
    results = {}
    for category, keywords in categories.items():
        for keyword in keywords:
            hits = search_word_paras(paras,keyword,whitelist)
            if (len(hits) > 0):
                res = { keyword: hits}
                if category not in results:
                    results[category]   = []
                results[category].append(res)
    return results

def get_pdf_paras(pdf_path: str) -> str:
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    parlist = text.split("\n")
    return [p for p in parlist if len(p) > 0]

def search_pdf_for_categories(pdf_path: str, categories:dict,whitelist:list) -> dict:

    paras  =get_pdf_paras(pdf_path)
    return search_paras_for_categories(paras,categories,whitelist)

def read_docs(directory: str) -> list:
    docx_files = glob.glob(os.path.join(directory, "*.docx"))
    doc_files = glob.glob(os.path.join(directory, "*.doc"))

    return docx_files + doc_files

def read_pdfs(directory: str) -> list:
    pdf_files = glob.glob(os.path.join(directory, "*.pdf"))
    return pdf_files

def check_files(docs: list, pdfs: list, categories: dict,whitelist: list) -> dict:
    res = {}
    for d in docs:
        try :
            doc = Document(d)
            r = search_doc_for_categories(doc,categories,whitelist)
            if len(r)>0:
                res[d]  = r
        except:
            print("unable to process..."+d)

    for p in pdfs: 
        ps = search_pdf_for_categories(p,categories,whitelist)
        if len(ps)>0:
            res[p] =ps
    return res


def main():
    # get args
    parser = argparse.ArgumentParser(description='Search for banned words in NIH proposal content.')
    parser.add_argument('--category-file', help='Path to the category JSON file')
    parser.add_argument('--directory', help='Path to the directory containing documents to search')
    parser.add_argument('--file',help='Single file to process')
    parser.add_argument('--whitelist',help='Whitelist file containing words that should be explicitly included')
    args = parser.parse_args()

    # 1. read in the json
    # 2. read in docs and pdfs.
    # routine to do all of them together...
    if args.category_file is None:
        print("Please provide a  JSON file with the banned categoriesusing --category-file")
        return
    
    if not os.path.isfile(args.category_file):
        print("category file "+args.category_file+" cannot be read.")
        return
    
    if args.directory is None and args.file is None:
        print("Please provide a directory containing documents to search using --directory or a file to process using --file")
        return
    categories = None
    with open(args.category_file, 'r') as f:
        categories = json.load(f)

    whitelist = None
    if args.whitelist is not None:
        if not os.path.isfile(args.whitelist):
            print("Whitelist file "+args.whitelist+" cannot be read")
            return
        with open(args.whitelist,'r') as f:
            whitelist = json.load(f)
            whitelist = [w.lower() for w in whitelist]

    if args.directory is not None:
        if not os.path.isdir(args.directory):
            print("Directory "+args.direcotry+" cannot be read")
            return
        docs = read_docs(args.directory)
        pdfs = read_pdfs(args.directory)
        results = check_files(docs,pdfs,categories,whitelist)
    else:
        results = {}
        if not os.path.isfile(args.file):
            print("File "+args.file+" cannot be read")
            return
        if args.file.endswith('.pdf'):
            results[args.file] = search_pdf_for_categories(args.file,categories,whitelist)
        else:
            doc = Document(args.file)
            results[args.file] = search_doc_for_categories(doc,categories,whitelist)
    print (json.dumps(results, indent=4))


if __name__ == "__main__":
    main()
